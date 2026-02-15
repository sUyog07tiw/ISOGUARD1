"""
Text extraction utilities for PDF, DOCX, and TXT files.
"""

import io
import logging
import re
from pathlib import Path
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


def sanitize_text(text: str) -> str:
    """
    Remove invalid Unicode characters (surrogates) from text.
    This is necessary because PDF extraction can produce malformed Unicode.
    
    Args:
        text: Input text that may contain invalid characters
        
    Returns:
        Sanitized text safe for UTF-8 encoding
    """
    if not text:
        return ""
    
    # Remove surrogate characters (U+D800 to U+DFFF)
    # These are invalid in UTF-8 and cause encoding errors
    sanitized = text.encode('utf-8', errors='surrogatepass').decode('utf-8', errors='replace')
    
    # Replace any remaining problematic characters with a space
    sanitized = re.sub(r'[\ud800-\udfff]', '', sanitized)
    
    # Also replace other control characters that might cause issues
    sanitized = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', sanitized)
    
    return sanitized


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_content: Raw bytes of the PDF file
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    # Sanitize text to remove invalid Unicode characters
                    page_text = sanitize_text(page_text)
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")
                continue
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise TextExtractionError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a DOCX file, preserving paragraph structure.
    
    Args:
        file_content: Raw bytes of the DOCX file
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Check if it's a heading
                if paragraph.style.name.startswith("Heading"):
                    text_parts.append(f"\n## {text}\n")
                else:
                    text_parts.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise TextExtractionError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from a TXT file.
    
    Args:
        file_content: Raw bytes of the TXT file
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If extraction fails
    """
    # Try common encodings
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252", "ascii"]
    
    for encoding in encodings:
        try:
            return file_content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    
    raise TextExtractionError("Failed to decode TXT file with any known encoding")


def extract_text(file_content: bytes, file_type: str) -> str:
    """
    Extract text from a file based on its type.
    
    Args:
        file_content: Raw bytes of the file
        file_type: Type of the file ('pdf', 'docx', 'txt')
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If extraction fails
        ValueError: If file type is unsupported
    """
    file_type = file_type.lower()
    
    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "txt": extract_text_from_txt,
    }
    
    if file_type not in extractors:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    return extractors[file_type](file_content)


def get_file_type(filename: str) -> Optional[str]:
    """
    Determine file type from filename extension.
    
    Args:
        filename: Name of the file
        
    Returns:
        File type string or None if unsupported
    """
    extension = Path(filename).suffix.lower().lstrip(".")
    
    # Only PDF is supported for audit reports
    supported_types = {"pdf"}
    
    if extension in supported_types:
        return extension
    
    return None
